from fastapi import FastAPI, File, UploadFile, HTTPException, Response, Request
from fastapi.middleware.cors import CORSMiddleware
import uuid
from datetime import datetime, timedelta
import asyncio

from typing import List, Dict, Any
from io import BytesIO, StringIO
from docx import Document
from langchain.docstore.document import Document as langchain_Document
from PyPDF2 import PdfReader
import csv

from langchain.prompts import ChatPromptTemplate, PromptTemplate
from langchain.prompts import SystemMessagePromptTemplate, HumanMessagePromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.memory import ConversationBufferMemory
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import Chroma
from langchain.chains import ConversationalRetrievalChain

from dotenv import load_dotenv

load_dotenv()


class Document_Processor:
    def __init__(self, file_details: List[Dict[Any, str]]):
        self.file_details = file_details

    def get_docs(self) -> List[langchain_Document]:
        docs = []
        for file_detail in self.file_details:
            if file_detail["name"].endswith(".txt"):
                docs.extend(self.get_txt_docs(file_detail=file_detail))

            elif file_detail["name"].endswith(".csv"):
                docs.extend(self.get_csv_docs(file_detail=file_detail))

            elif file_detail["name"].endswith(".docx"):
                docs.extend(self.get_docx_docs(file_detail=file_detail))

            elif file_detail["name"].endswith(".pdf"):
                docs.extend(self.get_pdf_docs(file_detail=file_detail))

        return docs

    @staticmethod
    def get_txt_docs(file_detail: Dict[str, Any]) -> List[langchain_Document]:
        text = file_detail["content"].decode("utf-8")
        source = file_detail["name"]
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000, chunk_overlap=100
        )
        text_docs = text_splitter.create_documents(
            [text], metadatas=[{"source": source}]
        )
        return text_docs

    @staticmethod
    def get_csv_docs(file_detail: Dict[str, Any]) -> List[langchain_Document]:
        csv_data = file_detail["content"]
        source = file_detail["name"]
        csv_string = csv_data.decode("utf-8")
        # Use StringIO to create a file-like object from the string
        csv_file = StringIO(csv_string)
        csv_reader = csv.DictReader(csv_file)
        csv_docs = []
        for row in csv_reader:
            # Convert each row into a dictionary of key/value pairs
            page_content = ""
            for key, value in row.items():
                page_content += f"{key}: {value}\n"
            doc = langchain_Document(
                page_content=page_content, metadata={"source": source}
            )
            csv_docs.append(doc)
        return csv_docs

    @staticmethod
    def get_pdf_docs(file_detail: Dict[str, Any]) -> List[langchain_Document]:
        pdf_content = BytesIO(file_detail["content"])
        source = file_detail["name"]

        reader = PdfReader(pdf_content)
        pdf_text = ""
        for page in reader.pages:
            pdf_text += page.extract_text() + "\n"

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000, chunk_overlap=100
        )
        pdf_docs = text_splitter.create_documents(
            texts=[pdf_text], metadatas=[{"source": source}]
        )
        return pdf_docs

    @staticmethod
    def get_docx_docs(file_detail: Dict[str, Any]) -> List[langchain_Document]:
        docx_content = BytesIO(file_detail["content"])
        source = file_detail["name"]

        document = Document(docx_content)
        docx_text = " ".join([paragraph.text for paragraph in document.paragraphs])

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000, chunk_overlap=100
        )
        docx_docs = text_splitter.create_documents(
            [docx_text], metadatas=[{"source": source}]
        )
        return docx_docs


class Conversational_Chain:
    def __init__(self, file_details: List[Dict[Any, str]]):
        self.llm_model = ChatOpenAI()
        self.embeddings = OpenAIEmbeddings()
        self.file_details = file_details

    def create_conversational_chain(self):
        docs = Document_Processor(self.file_details).get_docs()
        memory = ConversationBufferMemory(
            memory_key="chat_history", return_messages=True
        )
        vectordb = Chroma.from_documents(
            docs,
            self.embeddings,
        )
        retriever = vectordb.as_retriever()
        conversation_chain = ConversationalRetrievalChain.from_llm(
            llm=self.llm_model,
            retriever=retriever,
            condense_question_prompt=self.get_question_generator_prompt(),
            combine_docs_chain_kwargs={
                "document_prompt": self.get_document_prompt(),
                "prompt": self.get_final_prompt(),
            },
            memory=memory,
        )

        return conversation_chain

    @staticmethod
    def get_document_prompt() -> PromptTemplate:
        document_template = """Document Content:{page_content}
    Document Path: {source}"""
        return PromptTemplate(
            input_variables=["page_content", "source"],
            template=document_template,
        )

    @staticmethod
    def get_question_generator_prompt() -> PromptTemplate:
        question_generator_template = """Combine the chat history and follow up question into
    a standalone question.\n Chat History: {chat_history}\n
    Follow up question: {question}
    """
        return PromptTemplate.from_template(question_generator_template)

    @staticmethod
    def get_final_prompt() -> ChatPromptTemplate:
        final_prompt_template = """Answer question based on the context and chat_history.
    If you cannot find answers, ask more related questions from the user.
    Use only the basename of the file path as name of the documents.
    Mention document name of the documents you used in your answer.

    context:
    {context}

    chat_history:
    {chat_history}

    question:
    {question}

    Answer:
    """

        messages = [
            SystemMessagePromptTemplate.from_template(final_prompt_template),
            HumanMessagePromptTemplate.from_template("{question}"),
        ]

        return ChatPromptTemplate.from_messages(messages)


class UserSessionManager:
    def __init__(self):
        self.sessions = {}
        self.last_request_time = {}

    def get_session(self, user_id: str):
        if user_id not in self.sessions:
            self.sessions[user_id] = None
            self.last_request_time = datetime.now()
        return self.sessions[user_id]

    def set_session(self, user_id: str, conversational_chain):
        self.sessions[user_id] = conversational_chain
        self.last_request_time[user_id] = datetime.now()

    def delete_inactive_sessions(self, inactive_period: timedelta):
        current_time = datetime.now()
        for user_id, last_request_time in list(self.last_request_time.items()):
            if current_time - last_request_time > inactive_period:
                del self.sessions[user_id]
                del self.last_request_time[user_id]


app = FastAPI()

origins = ["https://viboognesh-react-chat.static.hf.space"]
# origins = ["http://localhost:3000"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["GET", "POST"],
    allow_headers=["*"],
)

user_session_manager = UserSessionManager()


@app.middleware("http")
async def update_last_request_time(request: Request, call_next):
    user_id = request.cookies.get("user_id")
    if user_id:
        user_session_manager.last_request_time[user_id] = datetime.now()
    response = await call_next(request)
    return response


async def check_inactivity():
    inactive_period = timedelta(hours=2)
    while True:
        await asyncio.sleep(600)
        user_session_manager.delete_inactive_sessions(inactive_period)


app.add_task(check_inactivity())


@app.post("/upload_files/")
async def upload_files(response: Response, files: List[UploadFile] = File(...)):
    file_details = []
    try:
        for file in files:
            content = await file.read()
            name = f"{file.filename}"
            details = {"content": content, "name": name}
            file_details.append(details)
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

    user_id = response.cookies.get("user_id")
    if not user_id:
        user_id = str(uuid.uuid4())
        response.set_cookie(key="user_id", value=user_id)

    try:
        conversational_chain = Conversational_Chain(
            file_details
        ).create_conversational_chain()
        user_session_manager.set_session(
            user_id=user_id, conversational_chain=conversational_chain
        )
        print("conversational_chain_manager created")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    return {"message": "ConversationalRetrievalChain is created. Please ask questions."}


@app.get("/predict/")
async def predict(query: str):
    user_id = response.cookies.get("user_id")
    if not user_id:
        user_id = str(uuid.uuid4())
        response.set_cookie(key="user_id", value=user_id)

    try:
        conversational_chain = user_session_manager.get_session(user_id=user_id)
        if conversational_chain is None:
            system_prompt = "Answer the question and also ask the user to upload files to ask questions from the files.\n"
            llm_model = ChatOpenAI()
            response = llm_model.invoke(system_prompt + query)
            answer = response.content
        else:
            response = conversational_chain.invoke(query)
            answer = response["answer"]
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    print("predict called")
    return {"answer": answer}
